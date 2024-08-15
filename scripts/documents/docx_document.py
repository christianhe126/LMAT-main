from documents.document import Document
import re
import base64
import sys
import json
from io import BytesIO
from PIL import Image
from docx import Document as _Document
from docx import opc
from docx import oxml
from docx.shared import Inches
from bs4 import BeautifulSoup
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Mm

class DocxDocument(Document):
    def __init__(self, format, data, template=None, metadata=None):
        super().__init__(format, data, template, metadata)
        self.document = _Document(self.path)
        self.document.add_page_break()
        if self.metadata:
            for key, value in self.metadata.items():
                self._replace_text("###" + key + "###", value)
        self.max_lines = 40
        self.chars_per_line = 90
        self.paragraphs = []
        self.page_size = {
            "x": 0,
            "y": 0,
            "width": Mm(210),
            "height": Mm(297)
        }

    def create_page(self, title = "[No Title]", following = None):
        if not following:
            self._add_title(self.document, title)
        return self.document

    def add_image(self, page, x, y, width, height, source):
        b64_string = source
        if b64_string.startswith("data:image/"):
            b64_string = b64_string.split(",")[1]
        
        image_data = base64.b64decode(b64_string)
        image_stream = BytesIO(image_data)
        x, y, width, height = self._calculate_img_size_and_position(x, y, width, height, image_data)
        width = width * 1.6
        height = height * 1.6

        if height > Inches(4):
            page.add_picture(image_stream, width=Inches(3))
        elif width > Inches(4):
            page.add_picture(image_stream, width=Inches(6))
        else:
            page.add_picture(image_stream, width=width)

        if hasattr(page, "paragraphs") and len(page.paragraphs) > 0:
            last_paragraph = page.paragraphs[-1] 
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_text(self, page, x, y, width, height, source):
        self.html_to_element(page, source)
    
    def add_table(self, page, x, y, width, height, source):
        n_row, n_col, b_style, b_width, b_color, cellspacing, data, sizes = self._get_table_data(source)
        shape = page.add_table(n_row, n_col)
        table = shape.table

        if b_style == "hidden" or b_style == "none" or b_width == 0 or b_width == "0":
            borders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                borders.append(border)
            table._tbl.tblPr.append(borders)
        else:
            table.style = 'Table Grid'
        
        for i in range(n_row):
            for j in range(n_col):
                if len(data) <= i or len(data[i]) <= j:
                    break
                cell = table.cell(i, j)
                # Check if the cell contains an image
                if re.compile(r'<[^>]+>').search(data[i][j]):
                    soup = BeautifulSoup(data[i][j], 'html.parser')
                    if soup.find("img") or soup.name == 'img' or (soup.name == 'figure' and soup.find('img')) or (soup.name == 'p' and soup.find('img')):
                        img = None
                        if hasattr(soup.find('img'), 'src'):
                            img = soup.find('img')
                        elif hasattr(soup, 'src'):
                            img = soup

                        if img:
                            src = img['src']
                            if src.startswith("data:image/"):
                                src = src.split(",")[1]
                            image_data = base64.b64decode(src)
                            with Image.open(BytesIO(image_data)) as _img:
                                img_width, img_height = _img.size
                            
                            ratio = img_height / img_width

                            width, height = table.cell(i, j).width, table.cell(i, j).width * ratio
                            if hasattr(img, 'width') and img.get("width") and hasattr(img, 'height') and img.get("height"):
                                x, y, width, height = self._calculate_img_size_and_position(x, y, int(img.get('width')), int(img.get('height')), image_data)
                            self.add_image(cell.paragraphs[0].add_run(), 0, 0, width*9600, height*9600, src)
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    else:
                        self.html_to_element(cell, data[i][j], cell.paragraphs[0])
                else:
                    self.html_to_element(cell, data[i][j], cell.paragraphs[0])

    def _reduce_paragraphs(self, paragraphs):
        i = 0
        while i < len(paragraphs) - 1:
            current = paragraphs[i]
            next_paragraph = paragraphs[i + 1]
            
            if current.type == "p" and next_paragraph.type == "p" and current.estimate + next_paragraph.estimate < self.max_lines:
                # Concatenate the contents and update the estimate
                current.content += next_paragraph.content
                current.estimate += next_paragraph.estimate
                # Remove the next paragraph from the list
                paragraphs.pop(i + 1)
            else:
                i += 1

    def _replace_text(self, old_text, new_text):
        # Replace text in paragraphs
        for paragraph in self.document.paragraphs:
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)
        
        # Replace text in tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if old_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(old_text, new_text)
        
        # Replace text in shapes
        for shape in self.document.inline_shapes:
            # Access the shape's XML element
            shape_element = shape._inline

            # Traverse the XML to find the text
            for paragraph in shape_element.xpath('.//w:p'):
                for run in paragraph.xpath('.//w:r'):
                    text_elements = run.xpath('.//w:t')
                    for text_element in text_elements:
                        if old_text in text_element.text:
                            # Replace the old text with the new text
                            text_element.text = text_element.text.replace(old_text, new_text)

    def _add_hyperlink_into_run(self, paragraph, run, url):
        runs = paragraph.runs
        for i in range(len(runs)):
            if runs[i].text == run.text:
                break

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(oxml.shared.qn('r:id'), r_id, )
        hyperlink.append(run._r)
        paragraph._p.insert(i+1,hyperlink)
    
    def _format_bullet_point(self, text):
        return chr(9679) + " " + text

    def _add_title(self, page, title):
        # Add a table with 1 row and 1 column
        table = self.document.add_table(rows=1, cols=1)

        # Get the cell
        cell = table.cell(0, 0)

        # Add the desired text to the cell
        paragraph = cell.paragraphs[0]
        if "---" not in title:
            run = paragraph.add_run(title)
            run.bold = True
        else:
            parts = title.split("---")
            for i, part in enumerate(parts):
                if i > 0:
                    paragraph.add_run("\n")
                run = paragraph.add_run(part)
                if i == 0:
                    run.bold = True
        
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Set font size and color
        font = run.font
        font.size = Pt(10)
        font.color.rgb = RGBColor(0, 0, 0)

        # Set background color
        shading_elm_1 = OxmlElement('w:shd')
        shading_elm_1.set(qn('w:fill'), 'D9D9D9')  # Light gray background
        cell._element.get_or_add_tcPr().append(shading_elm_1)

        # Set cell border
        cell_borders = cell._element.xpath('.//w:tcBorders')
        if not cell_borders:
            cell_borders = OxmlElement('w:tcBorders')
            cell._element.get_or_add_tcPr().append(cell_borders)

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Adjust the size as needed
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Black color
            cell_borders.append(border)

        self.add_text(page, 0, 0, 0, 0, "")


